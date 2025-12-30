#!/usr/bin/env python3
"""ë¬¸ì„œ ìƒì„± ëª¨ë“ˆ (Word/PDF)

JSON ë³´ê³ ì„œ ë°ì´í„°ë¥¼ Word/PDF ë¬¸ì„œë¡œ ë³€í™˜
"""

from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime
import re
import tempfile
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False


class DocumentGenerator:
    """Word/PDF ë¬¸ì„œ ìƒì„±ê¸°"""

    def __init__(self):
        self.first_heading_added = False
        # ì´ë¯¸ì§€ ê¸°ë³¸ ë””ë ‰í† ë¦¬ (í”„ë¡œì íŠ¸ ë£¨íŠ¸ ê¸°ì¤€)
        self.image_base_dir = Path(__file__).parent.parent / "data"

    def _generate_report_title(self, question: str) -> str:
        """ì§ˆë¬¸ ê¸°ë°˜ìœ¼ë¡œ ë³´ê³ ì„œ ì œëª© ìƒì„±

        Args:
            question: ì‚¬ìš©ì ì§ˆë¬¸

        Returns:
            ë³´ê³ ì„œ ì œëª©
        """
        # í‚¤ì›Œë“œ ê¸°ë°˜ ì œëª© ë§¤í•‘
        question_lower = question.lower()

        if "ìµœì¢…" in question or "ì¢…í•©" in question or "ì „ì²´" in question:
            return "AI/ML í”„ë¡œì íŠ¸ ìµœì¢… ë³´ê³ ì„œ"
        elif "ì£¼ê°„" in question or "weekly" in question_lower:
            return "ì£¼ê°„ ì—…ë¬´ ë³´ê³ ì„œ"
        elif "ì›”ê°„" in question or "monthly" in question_lower:
            return "ì›”ê°„ ì—…ë¬´ ë³´ê³ ì„œ"
        elif "ì„ì›" in question or "executive" in question_lower:
            return "ì„ì› ë³´ê³ ì„œ"
        elif "cmb" in question_lower or "ì¶”ì²œ" in question:
            return "CMB ì¶”ì²œì‹œìŠ¤í…œ ë³´ê³ ì„œ"
        elif "í…Œë‹ˆìŠ¤" in question or "ëª¨ë©˜í…€" in question:
            return "í…Œë‹ˆìŠ¤ ëª¨ë©˜í…€ ì˜ˆì¸¡ í”„ë¡œì íŠ¸ ë³´ê³ ì„œ"
        elif "ê¸‰ì´ëŸ‰" in question or "ì•„ì¿ ì•„" in question:
            return "ê¸‰ì´ëŸ‰ ë¶„ì„ í”„ë¡œì íŠ¸ ë³´ê³ ì„œ"
        elif "rag" in question_lower or "ì±—ë´‡" in question:
            return "RAG ì‹œìŠ¤í…œ êµ¬ì¶• ë³´ê³ ì„œ"
        else:
            # ê¸°ë³¸ ì œëª©
            return "í”„ë¡œì íŠ¸ ë³´ê³ ì„œ"

    def _should_include_image(self, image_info: Dict[str, Any], answer_text: str) -> bool:
        """ì´ë¯¸ì§€ë¥¼ ë³´ê³ ì„œì— í¬í•¨í• ì§€ ê²°ì •

        Args:
            image_info: ì´ë¯¸ì§€ ì •ë³´ ë”•ì…”ë„ˆë¦¬
            answer_text: ë‹µë³€ í…ìŠ¤íŠ¸

        Returns:
            í¬í•¨ ì—¬ë¶€ (True/False)
        """
        description = image_info.get('description', '').lower()
        source = image_info.get('source', '').lower()

        # ì œì™¸í•  ì´ë¯¸ì§€ íŒ¨í„´
        exclude_keywords = [
            'ì¼ì •', 'schedule', 'ê³„íší‘œ', 'í”„ë¡œì íŠ¸ ì¼ì •',
            'ì°¸ì„ì', 'participant', 'íšŒì˜ë¡',
            'ëª©ì°¨', 'table of contents'
        ]

        # ì„¤ëª…ì´ë‚˜ ì¶œì²˜ì— ì œì™¸ í‚¤ì›Œë“œê°€ ìˆìœ¼ë©´ ì œì™¸
        for keyword in exclude_keywords:
            if keyword in description or keyword in source:
                return False

        # í¬í•¨í•  ì´ë¯¸ì§€ íŒ¨í„´ (ì„±ëŠ¥, ê²°ê³¼, ë¶„ì„ ê´€ë ¨)
        include_keywords = [
            'ê²°ê³¼', 'result', 'ì„±ëŠ¥', 'performance',
            'ê·¸ë˜í”„', 'graph', 'ì°¨íŠ¸', 'chart',
            'ë¶„ì„', 'analysis', 'ì‹œê°í™”', 'visualization',
            'ëª¨ë¸', 'model', 'ì˜ˆì¸¡', 'prediction',
            'accuracy', 'precision', 'recall', 'f1',
            'ë¶„í¬', 'distribution', 'ë¹„êµ', 'comparison'
        ]

        # ì„¤ëª…ì´ë‚˜ ì¶œì²˜ì— í¬í•¨ í‚¤ì›Œë“œê°€ ìˆìœ¼ë©´ í¬í•¨
        for keyword in include_keywords:
            if keyword in description or keyword in source:
                return True

        # ê¸°ë³¸ì ìœ¼ë¡œ ì œì™¸
        return False

    def _shorten_image_caption(self, description: str, max_length: int = 100) -> str:
        """ì´ë¯¸ì§€ ìº¡ì…˜ì„ ì§§ê²Œ ìš”ì•½

        Args:
            description: ì›ë³¸ ì„¤ëª…
            max_length: ìµœëŒ€ ê¸¸ì´

        Returns:
            ìš”ì•½ëœ ì„¤ëª…
        """
        if not description or len(description) <= max_length:
            return description

        # ì²« ë¬¸ì¥ë§Œ ì¶”ì¶œ
        first_sentence = description.split('.')[0].split('ã€‚')[0]

        if len(first_sentence) <= max_length:
            return first_sentence

        # ê·¸ë˜ë„ ê¸¸ë©´ ìë¥´ê³  ... ì¶”ê°€
        return first_sentence[:max_length-3] + "..."

    def _add_image(self, doc: Document, image_path: str, description: str = None, max_width: float = 5.0):
        """ë¬¸ì„œì— ì´ë¯¸ì§€ ì¶”ê°€

        Args:
            doc: Document ê°ì²´
            image_path: ì´ë¯¸ì§€ ìƒëŒ€ ê²½ë¡œ (ì˜ˆ: "notion_images/xxx.png")
            description: ì´ë¯¸ì§€ ì„¤ëª… (ìº¡ì…˜)
            max_width: ìµœëŒ€ ë„ˆë¹„ (ì¸ì¹˜)
        """
        # ì ˆëŒ€ ê²½ë¡œë¡œ ë³€í™˜
        full_path = self.image_base_dir / image_path

        if not full_path.exists():
            print(f"âš ï¸ ì´ë¯¸ì§€ íŒŒì¼ì„ ì°¾ì„ ìˆ˜ ì—†ìŠµë‹ˆë‹¤: {full_path}")
            # ì´ë¯¸ì§€ê°€ ì—†ìœ¼ë©´ ì„¤ëª…ë§Œ í‘œì‹œ
            if description:
                para = doc.add_paragraph()
                run = para.add_run(f"[ì´ë¯¸ì§€: {description}]")
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
            return

        try:
            # ì´ë¯¸ì§€ ì¶”ê°€
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(str(full_path), width=Inches(max_width))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # ìº¡ì…˜ ì¶”ê°€
            if description:
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(description)
                caption_run.font.size = Pt(9)
                caption_run.font.italic = True
                caption_run.font.color.rgb = RGBColor(100, 100, 100)

        except Exception as e:
            print(f"âš ï¸ ì´ë¯¸ì§€ ì‚½ì… ì‹¤íŒ¨: {full_path}, ì˜¤ë¥˜: {e}")
            # ì‹¤íŒ¨ ì‹œ ì„¤ëª…ë§Œ í‘œì‹œ
            if description:
                para = doc.add_paragraph()
                run = para.add_run(f"[ì´ë¯¸ì§€: {description}]")
                run.font.size = Pt(10)
                run.font.italic = True

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """í—¤ë”© ì¶”ê°€"""
        heading = doc.add_heading(text, level=level)
        return heading

    def _add_paragraph(self, doc: Document, text: str, bold: bool = False, italic: bool = False):
        """ë‹¨ë½ ì¶”ê°€"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Malgun Gothic'

        if bold:
            run.bold = True
        if italic:
            run.italic = True

        return para

    def _parse_markdown_table(self, table_text: str) -> List[List[str]]:
        """ë§ˆí¬ë‹¤ìš´ í…Œì´ë¸” íŒŒì‹±

        Args:
            table_text: ë§ˆí¬ë‹¤ìš´ í…Œì´ë¸” í…ìŠ¤íŠ¸

        Returns:
            2D ë¦¬ìŠ¤íŠ¸ (í–‰xì—´)
        """
        lines = table_text.strip().split('\n')
        rows = []

        for i, line in enumerate(lines):
            # ë¹ˆ ì¤„ ìŠ¤í‚µ
            if not line.strip():
                continue

            # í…Œì´ë¸” í–‰ì´ ì•„ë‹ˆë©´ ìŠ¤í‚µ
            if '|' not in line:
                continue

            # êµ¬ë¶„ì„  ìŠ¤í‚µ (ì˜ˆ: |---|---|, | --- | --- |, |-----|-----|)
            stripped = line.strip()
            # í•˜ì´í”ˆìœ¼ë¡œë§Œ êµ¬ì„±ëœ ì…€ì´ ìˆìœ¼ë©´ êµ¬ë¶„ì„ ìœ¼ë¡œ ê°„ì£¼
            temp_cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if temp_cells and all(set(cell) <= set('-: ') for cell in temp_cells):
                continue

            # ì…€ ì¶”ì¶œ: | ê¸°ì¤€ìœ¼ë¡œ ë¶„ë¦¬
            parts = line.split('|')

            # ë§¨ ì•ë’¤ ë¹ˆ ë¬¸ìì—´ ì œê±° (|ë¡œ ì‹œì‘í•˜ê³  ëë‚˜ëŠ” ê²½ìš°)
            if parts and not parts[0].strip():
                parts = parts[1:]
            if parts and not parts[-1].strip():
                parts = parts[:-1]

            # ê° ì…€ì˜ ì•ë’¤ ê³µë°±ë§Œ ì œê±° (ë¹ˆ ì…€ ìœ ì§€)
            cells = [cell.strip() for cell in parts]

            if cells:
                rows.append(cells)

        return rows

    def _add_markdown_table(self, doc: Document, table_text: str):
        """ë§ˆí¬ë‹¤ìš´ í…Œì´ë¸”ì„ Word í…Œì´ë¸”ë¡œ ë³€í™˜

        Args:
            doc: Word ë¬¸ì„œ
            table_text: ë§ˆí¬ë‹¤ìš´ í…Œì´ë¸” í…ìŠ¤íŠ¸
        """
        rows_data = self._parse_markdown_table(table_text)

        if not rows_data or len(rows_data) == 0:
            print(f"âš ï¸ í…Œì´ë¸” íŒŒì‹± ì‹¤íŒ¨ ë˜ëŠ” ë¹ˆ í…Œì´ë¸”")
            return

        # ëª¨ë“  í–‰ì˜ ì—´ ê°œìˆ˜ í™•ì¸ (ê°€ì¥ ë§ì€ ì—´ì„ ê¸°ì¤€ìœ¼ë¡œ)
        num_cols = max(len(row) for row in rows_data)
        num_rows = len(rows_data)

        # ê° í–‰ì˜ ì—´ ê°œìˆ˜ë¥¼ ë§ì¶¤ (ë¶€ì¡±í•œ ê²½ìš° ë¹ˆ ì…€ ì¶”ê°€)
        for row in rows_data:
            while len(row) < num_cols:
                row.append('')

        print(f"ğŸ“Š í…Œì´ë¸” ìƒì„±: {num_rows}í–‰ x {num_cols}ì—´")

        # Word í…Œì´ë¸” ìƒì„±
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # ë°ì´í„° ì±„ìš°ê¸°
        for i, row_data in enumerate(rows_data):
            for j, cell_data in enumerate(row_data):
                if j < num_cols and i < num_rows:
                    cell = table.rows[i].cells[j]
                    cell.text = str(cell_data) if cell_data else ''

                    # ì²« í–‰ì€ í—¤ë”ë¡œ ë³¼ë“œ ì²˜ë¦¬
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(10)
                                run.font.name = 'Malgun Gothic'
                    else:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                                run.font.name = 'Malgun Gothic'

        doc.add_paragraph()  # í…Œì´ë¸” ë’¤ ê°„ê²©

    def _add_formatted_text(self, doc: Document, text: str):
        """ë§ˆí¬ë‹¤ìš´ í˜•ì‹ì´ í¬í•¨ëœ í…ìŠ¤íŠ¸ë¥¼ Wordì— ì¶”ê°€

        Args:
            doc: Word ë¬¸ì„œ
            text: ë§ˆí¬ë‹¤ìš´ í˜•ì‹ í…ìŠ¤íŠ¸
        """
        lines = text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # í…Œì´ë¸” ì‹œì‘ ê°ì§€: | ë¡œ ì‹œì‘í•˜ê³  ëë‚˜ëŠ” í–‰
            if line.strip().startswith('|') and '|' in line:
                # í…Œì´ë¸” ë¸”ë¡ ìˆ˜ì§‘
                table_lines = []
                start_i = i

                # ì—°ì†ëœ í…Œì´ë¸” í–‰ ìˆ˜ì§‘
                while i < len(lines):
                    current_line = lines[i].strip()

                    # ë¦¬ìŠ¤íŠ¸ ë§ˆì»¤ê°€ ìˆëŠ” ê²½ìš° ì œê±°
                    current_line = re.sub(r'^\s*[-*+]\s+', '', current_line)
                    current_line = re.sub(r'^\s+', '', current_line)

                    # í…Œì´ë¸” í–‰ì¸ì§€ í™•ì¸
                    if current_line.startswith('|') and '|' in current_line:
                        table_lines.append(current_line)
                        i += 1
                    else:
                        # í…Œì´ë¸”ì´ ëë‚¨
                        break

                # ìˆ˜ì§‘í•œ í…Œì´ë¸” ë¸”ë¡ ì²˜ë¦¬
                if table_lines:
                    table_text = '\n'.join(table_lines)
                    print(f"ğŸ” í…Œì´ë¸” ê°ì§€ ({start_i}í–‰ë¶€í„° {len(table_lines)}ì¤„):")
                    print(table_text[:200] + '...' if len(table_text) > 200 else table_text)
                    self._add_markdown_table(doc, table_text)
            else:
                # ì¼ë°˜ í…ìŠ¤íŠ¸ ì²˜ë¦¬ - ë‹¤ìŒ í…Œì´ë¸”ê¹Œì§€ì˜ ëª¨ë“  í–‰ ìˆ˜ì§‘
                text_lines = []
                while i < len(lines):
                    if lines[i].strip().startswith('|') and '|' in lines[i]:
                        # ë‹¤ìŒ í…Œì´ë¸” ë°œê²¬
                        break
                    text_lines.append(lines[i])
                    i += 1

                if text_lines:
                    paragraph_text = '\n'.join(text_lines)
                    self._add_formatted_paragraph(doc, paragraph_text)

    def _add_formatted_paragraph(self, doc: Document, text: str):
        """ë§ˆí¬ë‹¤ìš´ í˜•ì‹(ë³¼ë“œ, ì´íƒ¤ë¦­, ë¦¬ìŠ¤íŠ¸)ì„ Wordë¡œ ë³€í™˜

        Args:
            doc: Word ë¬¸ì„œ
            text: í…ìŠ¤íŠ¸
        """
        lines = text.strip().split('\n')

        for line in lines:
            if not line.strip():
                continue

            # ìˆ˜í‰ì„  ê°ì§€ (---, ___, ***)
            if re.match(r'^[\-_*]{3,}\s*$', line):
                # ìˆ˜í‰ì„ ì€ ìŠ¤í‚µ (í‘œì‹œí•˜ì§€ ì•ŠìŒ)
                continue

            # ë¦¬ìŠ¤íŠ¸ í•­ëª© ê°ì§€
            list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
            if list_match:
                indent = len(list_match.group(1))
                content = list_match.group(3)
                para = doc.add_paragraph(style='List Bullet')
                self._add_inline_formatting(para, content)
                continue

            # í—¤ë”© ê°ì§€
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                content = heading_match.group(2)
                heading = doc.add_heading(content, level=level)

                # ì²« ë²ˆì§¸ ë ˆë²¨ 1 í—¤ë”©ë§Œ ì¤‘ì•™ ì •ë ¬
                if not self.first_heading_added and level == 1:
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.first_heading_added = True

                continue

            # ì¼ë°˜ ë‹¨ë½
            para = doc.add_paragraph()
            self._add_inline_formatting(para, line)

    def _add_inline_formatting(self, paragraph, text: str):
        """ì¸ë¼ì¸ ë§ˆí¬ë‹¤ìš´ í˜•ì‹(ë³¼ë“œ, ì´íƒ¤ë¦­) ì²˜ë¦¬

        Args:
            paragraph: Word ë‹¨ë½
            text: í…ìŠ¤íŠ¸
        """
        # ë³¼ë“œ+ì´íƒ¤ë¦­: ***text*** ë˜ëŠ” ___text___
        # ë³¼ë“œ: **text** ë˜ëŠ” __text__
        # ì´íƒ¤ë¦­: *text* ë˜ëŠ” _text_

        pattern = r'(\*\*\*|___|__|\*\*|_|\*)(.*?)\1'
        last_pos = 0

        for match in re.finditer(pattern, text):
            # ì¼ë°˜ í…ìŠ¤íŠ¸ ì¶”ê°€
            if match.start() > last_pos:
                run = paragraph.add_run(text[last_pos:match.start()])
                run.font.size = Pt(11)
                run.font.name = 'Malgun Gothic'

            # í˜•ì‹í™”ëœ í…ìŠ¤íŠ¸ ì¶”ê°€
            marker = match.group(1)
            content = match.group(2)

            run = paragraph.add_run(content)
            run.font.size = Pt(11)
            run.font.name = 'Malgun Gothic'

            if marker in ['***', '___']:
                run.bold = True
                run.italic = True
            elif marker in ['**', '__']:
                run.bold = True
            elif marker in ['*', '_']:
                run.italic = True

            last_pos = match.end()

        # ë‚¨ì€ í…ìŠ¤íŠ¸ ì¶”ê°€
        if last_pos < len(text):
            run = paragraph.add_run(text[last_pos:])
            run.font.size = Pt(11)
            run.font.name = 'Malgun Gothic'

    def generate_word_report(self, report_data: Dict[str, Any], output_path: str):
        """Word ë³´ê³ ì„œ ìƒì„±

        Args:
            report_data: ë³´ê³ ì„œ ë°ì´í„° (JSON)
            output_path: ì¶œë ¥ íŒŒì¼ ê²½ë¡œ
        """
        # Pandoc ì‚¬ìš© ê°€ëŠ¥ ì—¬ë¶€ í™•ì¸
        if PANDOC_AVAILABLE:
            try:
                self._generate_word_with_pandoc_and_tables(report_data, output_path)
                return
            except Exception as e:
                print(f"âš ï¸ Pandoc ë³€í™˜ ì‹¤íŒ¨, ê¸°ë³¸ ë°©ì‹ìœ¼ë¡œ ì „í™˜: {e}")

        # ê¸°ë³¸ ë°©ì‹ (python-docx)
        print("ğŸ”§ python-docx ë°©ì‹ìœ¼ë¡œ Word ìƒì„± (í…Œì´ë¸” ì§€ì›)")
        self._generate_word_basic(report_data, output_path)

    def _remove_first_heading(self, text: str) -> str:
        """ë¶ˆí•„ìš”í•œ í—¤ë”© ì œê±° ë° ì •ë¦¬

        Args:
            text: ë§ˆí¬ë‹¤ìš´ í…ìŠ¤íŠ¸

        Returns:
            ìˆ˜ì •ëœ í…ìŠ¤íŠ¸
        """
        import re
        lines = text.split('\n')
        result_lines = []
        skip_mode = False  # Executive Summary ì„¹ì…˜ ì „ì²´ ìŠ¤í‚µìš©

        for line in lines:
            stripped = line.strip()

            # "ì„ì› ë³´ê³ ì„œ", "ìµœì¢… ë³´ê³ ì„œ" ë“± ë¶ˆí•„ìš”í•œ í—¤ë”© ì œê±°
            if re.match(r'^#{1,3}\s*(ì„ì›\s*ë³´ê³ ì„œ|ìµœì¢…\s*ë³´ê³ ì„œ|ì¼ì£¼ì¼\s*ë³´ê³ ì„œ|ì£¼ê°„\s*ë³´ê³ ì„œ)', stripped):
                continue

            # "Executive Summary (í•µì‹¬ ìš”ì•½)" ì„¹ì…˜ ì‹œì‘ - ë‹¤ìŒ ## í—¤ë”©ê¹Œì§€ ìŠ¤í‚µ
            if re.match(r'^#{1,3}\s*Executive Summary', stripped, re.IGNORECASE):
                skip_mode = True
                continue

            # skip_mode ì¤‘ ë‹¤ìŒ ì„¹ì…˜ í—¤ë”© ë°œê²¬ ì‹œ ìŠ¤í‚µ ëª¨ë“œ ì¢…ë£Œ
            if skip_mode:
                # ## ë˜ëŠ” ### ë¡œ ì‹œì‘í•˜ëŠ” ë‹¤ë¥¸ ì„¹ì…˜ì´ ë‚˜ì˜¤ë©´ ìŠ¤í‚µ ëª¨ë“œ ì¢…ë£Œ
                if re.match(r'^#{1,3}\s*\d+\.', stripped):
                    skip_mode = False
                    # í˜„ì¬ ë¼ì¸ ì²˜ë¦¬ë¡œ ë„˜ì–´ê° (continue í•˜ì§€ ì•ŠìŒ)
                else:
                    # Executive Summary ì„¹ì…˜ ë‚´ìš© ìŠ¤í‚µ
                    continue

            # ì„¹ì…˜ ì œëª©ì„ ## (ë ˆë²¨ 2)ë¡œ í†µì¼
            # "### 1. ê²°ê³¼" â†’ "## 1. ê²°ê³¼"
            if re.match(r'^#{1,4}\s*1\.', stripped):
                # ìˆ«ì ë‹¤ìŒ í…ìŠ¤íŠ¸ ì¶”ì¶œ
                section_text = re.sub(r'^#{1,4}\s*', '', stripped)
                result_lines.append(f'## {section_text}')
                continue

            # "### 2. ì£¼ìš” í˜„í™©" â†’ "## 2. ì£¼ìš” í˜„í™©"
            if re.match(r'^#{1,4}\s*2\.', stripped):
                section_text = re.sub(r'^#{1,4}\s*', '', stripped)
                result_lines.append(f'## {section_text}')
                continue

            # "### 3. í•µì‹¬ ì´ìŠˆ ë° ë¦¬ìŠ¤í¬" â†’ "## 3. í•µì‹¬ ì´ìŠˆ ë° ë¦¬ìŠ¤í¬"
            if re.match(r'^#{1,4}\s*3\.', stripped):
                section_text = re.sub(r'^#{1,4}\s*', '', stripped)
                result_lines.append(f'## {section_text}')
                continue

            # "### 4. ì¶”ê°€ í™•ì¸ í•„ìš” ì‚¬í•­" â†’ "## 4. ì¶”ê°€ í™•ì¸ í•„ìš” ì‚¬í•­"
            if re.match(r'^#{1,4}\s*4\.', stripped):
                section_text = re.sub(r'^#{1,4}\s*', '', stripped)
                result_lines.append(f'## {section_text}')
                continue

            result_lines.append(line)

        return '\n'.join(result_lines)

    def _fix_table_format(self, text: str) -> str:
        """ìœ ë‹ˆì½”ë“œ ë°•ìŠ¤ ë¬¸ìë¥¼ ë§ˆí¬ë‹¤ìš´ í…Œì´ë¸”ë¡œ ë³€í™˜í•˜ê³  êµ¬ë¶„ì„  ì¶”ê°€

        ë¦¬ìŠ¤íŠ¸ ì•ˆì˜ í…Œì´ë¸”ì„ ë…ë¦½ì ì¸ í…Œì´ë¸” ë¸”ë¡ìœ¼ë¡œ ë³€í™˜
        """
        # â”‚ (ìœ ë‹ˆì½”ë“œ ë°•ìŠ¤ ë¬¸ì)ë¥¼ | (íŒŒì´í”„)ë¡œ ë³€í™˜
        text = text.replace('â”‚', '|')
        # â”€ (ìœ ë‹ˆì½”ë“œ ê°€ë¡œì„ )ë¥¼ - (í•˜ì´í”ˆ)ë¡œ ë³€í™˜
        text = text.replace('â”€', '-')

        lines = text.split('\n')
        fixed_lines = []
        in_table = False
        table_buffer = []

        for i, line in enumerate(lines):
            stripped = line.strip()

            # í…Œì´ë¸” í–‰ì¸ì§€ í™•ì¸ (ë¦¬ìŠ¤íŠ¸ ë§ˆì»¤ í¬í•¨)
            is_table_line = False
            clean_line = line

            # ë¦¬ìŠ¤íŠ¸ ë§ˆì»¤ë¡œ ì‹œì‘í•˜ëŠ” í…Œì´ë¸” í–‰
            if re.match(r'^\s*[-*+]\s+\|', line):
                is_table_line = True
                clean_line = re.sub(r'^\s*[-*+]\s+', '', line)
            # ë“¤ì—¬ì“°ê¸°ëœ í…Œì´ë¸” í–‰
            elif re.match(r'^\s+\|', line):
                is_table_line = True
                clean_line = stripped
            # ì¼ë°˜ í…Œì´ë¸” í–‰
            elif stripped.startswith('|') and stripped.endswith('|'):
                is_table_line = True
                clean_line = stripped

            if is_table_line:
                if not in_table:
                    in_table = True
                    # í…Œì´ë¸” ì‹œì‘ ì „ì— ë¹ˆ ì¤„ ì¶”ê°€
                    if fixed_lines and fixed_lines[-1].strip():
                        fixed_lines.append('')
                table_buffer.append(clean_line)
            else:
                # í…Œì´ë¸”ì´ ëë‚¬ìœ¼ë©´ ë²„í¼ ì²˜ë¦¬
                if in_table:
                    # í…Œì´ë¸” ì •ë¦¬ ë° ì¶”ê°€
                    self._finalize_table_buffer(table_buffer, fixed_lines)
                    table_buffer = []
                    in_table = False
                    # í…Œì´ë¸” ë’¤ì— ë¹ˆ ì¤„ ì¶”ê°€
                    fixed_lines.append('')

                fixed_lines.append(line)

        # ë§ˆì§€ë§‰ í…Œì´ë¸” ì²˜ë¦¬
        if in_table and table_buffer:
            if fixed_lines and fixed_lines[-1].strip():
                fixed_lines.append('')
            self._finalize_table_buffer(table_buffer, fixed_lines)
            fixed_lines.append('')

        return '\n'.join(fixed_lines)

    def _finalize_table_buffer(self, table_buffer: list, output_lines: list):
        """í…Œì´ë¸” ë²„í¼ë¥¼ ì •ë¦¬í•˜ê³  ì¶œë ¥ ë¼ì¸ì— ì¶”ê°€"""
        if not table_buffer:
            return

        # ì²« ë²ˆì§¸ í–‰ì´ í—¤ë”
        header = table_buffer[0]

        # ë‘ ë²ˆì§¸ í–‰ì´ êµ¬ë¶„ì„ ì¸ì§€ í™•ì¸
        has_separator = False
        separator_idx = -1
        data_start_idx = 1

        # ë²„í¼ì—ì„œ êµ¬ë¶„ì„  ì°¾ê¸° (ì²« ëª‡ í–‰ì—ì„œë§Œ)
        for idx in range(1, min(3, len(table_buffer))):
            line = table_buffer[idx].strip()
            # êµ¬ë¶„ì„  íŒ¨í„´: |-----|-----|  ë˜ëŠ” | --- | --- |
            if re.match(r'^\|[\s\-:|]+\|$', line) and '-' in line:
                has_separator = True
                separator_idx = idx
                data_start_idx = idx + 1
                break

        # í—¤ë” ì¶”ê°€
        output_lines.append(header)

        # êµ¬ë¶„ì„  ì¶”ê°€ (ìˆìœ¼ë©´ ì›ë³¸ ì‚¬ìš©, ì—†ìœ¼ë©´ ìƒì„±)
        if has_separator:
            output_lines.append(table_buffer[separator_idx])
        else:
            num_cols = header.count('|') - 1
            separator = '|' + '|'.join(['---' for _ in range(num_cols)]) + '|'
            output_lines.append(separator)

        # ë‚˜ë¨¸ì§€ ë°ì´í„° í–‰ ì¶”ê°€
        for row in table_buffer[data_start_idx:]:
            output_lines.append(row)

    def _generate_word_with_pandoc_and_tables(self, report_data: Dict[str, Any], output_path: str):
        """Pandoc + python-docx í•˜ì´ë¸Œë¦¬ë“œ ë°©ì‹ìœ¼ë¡œ Word ìƒì„±

        í‘œëŠ” python-docxë¡œ ì§ì ‘ ìƒì„±í•˜ê³ , ë‚˜ë¨¸ì§€ëŠ” pandocìœ¼ë¡œ ë³€í™˜
        í‘œëŠ” ì›ë˜ ë§ˆí¬ë‹¤ìš´ì— ìˆë˜ ìœ„ì¹˜ì— ì •í™•íˆ ë°°ì¹˜
        """
        print("ğŸ”§ í•˜ì´ë¸Œë¦¬ë“œ ë°©ì‹: í‘œëŠ” python-docx, ë‚˜ë¨¸ì§€ëŠ” Pandoc")

        # ê²°ê³¼ ìˆ˜ì§‘ ë° í‘œ ì¶”ì¶œ
        results = report_data.get('results', [])
        markdown_content = []
        all_tables = []  # ëª¨ë“  í‘œë¥¼ ìˆœì„œëŒ€ë¡œ ì €ì¥

        # ë³´ê³ ì„œ ì œëª© ì¶”ê°€ (ì§ˆë¬¸ ê¸°ë°˜)
        if results and results[0].get('question'):
            title = self._generate_report_title(results[0]['question'])
            markdown_content.append(f'# {title}\n\n')

            # ì‘ì„±ì ë° ì‘ì„±ì¼ì ì¶”ê°€
            author = report_data.get('author', 'Unknown')
            created_date = report_data.get('created_date', datetime.now().strftime("%Y-%m-%d"))
            markdown_content.append(f'**ì‘ì„±ì:** {author}  |  **ì‘ì„±ì¼:** {created_date}\n\n')
            markdown_content.append('---\n\n')

        # ê° ë‹µë³€ì„ ì²˜ë¦¬í•˜ë©° í‘œë¥¼ placeholderë¡œ ì¹˜í™˜
        for result in results:
            if result.get('success'):
                answer = result.get('answer', 'N/A')
                # í…Œì´ë¸” í˜•ì‹ ìˆ˜ì •
                answer = self._fix_table_format(answer)
                # "ì„ì›ë³´ê³ ì„œ" í—¤ë”© ì œê±°
                answer = self._remove_first_heading(answer)

                # í‘œë¥¼ placeholderë¡œ ì¹˜í™˜
                answer_with_placeholders, tables = self._replace_tables_with_placeholders(answer, len(all_tables))
                all_tables.extend(tables)

                markdown_content.append(answer_with_placeholders)
                markdown_content.append('\n\n')

        full_markdown = '\n'.join(markdown_content)

        # ì„ì‹œ ë§ˆí¬ë‹¤ìš´ íŒŒì¼ ìƒì„±
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp:
            tmp.write(full_markdown)
            tmp_path = tmp.name

        print(f"ğŸ“ ì„ì‹œ ë§ˆí¬ë‹¤ìš´ íŒŒì¼: {tmp_path}")
        if all_tables:
            print(f"ğŸ“Š ì´ {len(all_tables)}ê°œì˜ í‘œ ë°œê²¬")

        try:
            # Pandocìœ¼ë¡œ ë³€í™˜
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)

            pypandoc.convert_file(
                tmp_path,
                'docx',
                outputfile=str(output_file),
                extra_args=['--reference-doc='] if False else []
            )

            print(f"ğŸ“„ Pandoc ë³€í™˜ ì™„ë£Œ, í‘œ ì‚½ì… ì¤‘...")

            # Word ë¬¸ì„œ ì—´ê¸°
            doc = Document(output_path)

            # Placeholderë¥¼ ì‹¤ì œ í‘œë¡œ êµì²´
            self._replace_placeholders_with_tables(doc, all_tables)

            # ìˆ˜ì •ëœ ë¬¸ì„œ ì €ì¥
            doc.save(output_path)
            print(f"âœ… Word ë¬¸ì„œ ìƒì„± ì™„ë£Œ: {output_path}")

        finally:
            # ì„ì‹œ íŒŒì¼ ì‚­ì œ
            Path(tmp_path).unlink(missing_ok=True)

    def _replace_tables_with_placeholders(self, text: str, start_index: int = 0) -> tuple:
        """ë§ˆí¬ë‹¤ìš´ í…ìŠ¤íŠ¸ì—ì„œ í‘œë¥¼ placeholderë¡œ ì¹˜í™˜

        Args:
            text: ë§ˆí¬ë‹¤ìš´ í…ìŠ¤íŠ¸
            start_index: í‘œ ë²ˆí˜¸ ì‹œì‘ ì¸ë±ìŠ¤

        Returns:
            (ì¹˜í™˜ëœ í…ìŠ¤íŠ¸, ì¶”ì¶œëœ í‘œ ë¦¬ìŠ¤íŠ¸)
        """
        lines = text.split('\n')
        result_lines = []
        tables = []
        i = 0
        table_index = start_index

        while i < len(lines):
            line = lines[i]

            # í…Œì´ë¸” ì‹œì‘ ê°ì§€
            if line.strip().startswith('|') and '|' in line:
                table_lines = []

                # ì—°ì†ëœ í…Œì´ë¸” í–‰ ìˆ˜ì§‘
                while i < len(lines):
                    current_line = lines[i].strip()
                    if current_line.startswith('|') and '|' in current_line:
                        table_lines.append(current_line)
                        i += 1
                    else:
                        break

                if table_lines:
                    # í‘œë¥¼ ì €ì¥í•˜ê³  placeholder ì‚½ì…
                    table_text = '\n'.join(table_lines)
                    tables.append(table_text)
                    result_lines.append(f'[TABLE_{table_index}]')
                    table_index += 1
            else:
                result_lines.append(line)
                i += 1

        return '\n'.join(result_lines), tables

    def _extract_tables_from_markdown(self, text: str) -> List[str]:
        """ë§ˆí¬ë‹¤ìš´ í…ìŠ¤íŠ¸ì—ì„œ í…Œì´ë¸” ì¶”ì¶œ"""
        tables = []
        lines = text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # í…Œì´ë¸” ì‹œì‘ ê°ì§€
            if line.strip().startswith('|') and '|' in line:
                table_lines = []

                # ì—°ì†ëœ í…Œì´ë¸” í–‰ ìˆ˜ì§‘
                while i < len(lines):
                    current_line = lines[i].strip()
                    if current_line.startswith('|') and '|' in current_line:
                        table_lines.append(current_line)
                        i += 1
                    else:
                        break

                if table_lines:
                    tables.append('\n'.join(table_lines))
            else:
                i += 1

        return tables

    def _replace_placeholders_with_tables(self, doc: Document, markdown_tables: List[str]):
        """Word ë¬¸ì„œì˜ placeholderë¥¼ python-docx í‘œë¡œ êµì²´

        Args:
            doc: Word ë¬¸ì„œ ê°ì²´
            markdown_tables: ë§ˆí¬ë‹¤ìš´ í‘œ ë¦¬ìŠ¤íŠ¸ (ìˆœì„œëŒ€ë¡œ)
        """
        from docx.table import Table

        table_idx = 0

        # ëª¨ë“  ë‹¨ë½ì„ ìˆœíšŒí•˜ë©° placeholder ì°¾ê¸°
        for para_idx in range(len(doc.paragraphs)):
            paragraph = doc.paragraphs[para_idx]
            text = paragraph.text.strip()

            # Placeholder íŒ¨í„´ í™•ì¸
            if text.startswith('[TABLE_') and text.endswith(']'):
                # í‘œ ë²ˆí˜¸ ì¶”ì¶œ
                try:
                    placeholder_num = int(text[7:-1])  # [TABLE_X]ì—ì„œ X ì¶”ì¶œ
                except:
                    continue

                if placeholder_num >= len(markdown_tables):
                    continue

                md_table = markdown_tables[placeholder_num]

                # í…Œì´ë¸” íŒŒì‹±
                rows_data = self._parse_markdown_table(md_table)
                if not rows_data:
                    continue

                num_rows = len(rows_data)
                num_cols = max(len(row) for row in rows_data)

                # ê° í–‰ì˜ ì—´ ê°œìˆ˜ë¥¼ ë§ì¶¤
                for row in rows_data:
                    while len(row) < num_cols:
                        row.append('')

                # Placeholder ë‹¨ë½ ìœ„ì¹˜ì— í…Œì´ë¸” ì‚½ì…
                p_element = paragraph._element
                parent = p_element.getparent()

                # ìƒˆ í…Œì´ë¸” ìƒì„±
                tbl = doc.add_table(rows=num_rows, cols=num_cols)._element

                # Placeholder ë‹¨ë½ ë°”ë¡œ ì•ì— í…Œì´ë¸” ì‚½ì…
                parent.insert(parent.index(p_element), tbl)

                # í…Œì´ë¸” ê°ì²´ ê°€ì ¸ì˜¤ê¸° ë° ìŠ¤íƒ€ì¼ ì„¤ì •
                new_table = Table(tbl, doc)

                try:
                    new_table.style = 'Light Grid Accent 1'
                except:
                    try:
                        new_table.style = 'Table Grid'
                    except:
                        pass

                # ë°ì´í„° ì±„ìš°ê¸°
                for i, row_data in enumerate(rows_data):
                    for j, cell_data in enumerate(row_data):
                        if j < num_cols and i < num_rows:
                            cell = new_table.rows[i].cells[j]
                            cell.text = str(cell_data) if cell_data else ''

                            # ì²« í–‰ì€ í—¤ë”ë¡œ ë³¼ë“œ ì²˜ë¦¬
                            if i == 0:
                                for cell_para in cell.paragraphs:
                                    for run in cell_para.runs:
                                        run.bold = True
                                        run.font.size = Pt(10)
                                        run.font.name = 'Malgun Gothic'
                            else:
                                for cell_para in cell.paragraphs:
                                    for run in cell_para.runs:
                                        run.font.size = Pt(9)
                                        run.font.name = 'Malgun Gothic'

                # Placeholder ë‹¨ë½ ì‚­ì œ
                p_element.getparent().remove(p_element)

                print(f"âœ… í‘œ {placeholder_num} ì‚½ì… ì™„ë£Œ ({num_rows}í–‰ x {num_cols}ì—´)")

    def _replace_tables_in_word(self, doc: Document, markdown_tables: List[str]):
        """Word ë¬¸ì„œì˜ í…Œì´ë¸”ì„ python-docxë¡œ ì¬ìƒì„±í•œ í…Œì´ë¸”ë¡œ êµì²´"""
        from docx.oxml import OxmlElement

        # í…ìŠ¤íŠ¸ë¡œ ë Œë”ë§ëœ í…Œì´ë¸” í–‰ë“¤ì„ ì°¾ì•„ì„œ ì‚­ì œí•˜ê³  ê·¸ ìë¦¬ì— ì‹¤ì œ í…Œì´ë¸” ì‚½ì…
        paragraphs_to_remove = []
        table_insert_positions = []

        i = 0
        while i < len(doc.paragraphs):
            paragraph = doc.paragraphs[i]
            text = paragraph.text.strip()

            # í…Œì´ë¸” ì‹œì‘ ê°ì§€ (| ë¡œ ì‹œì‘í•˜ëŠ” ì¤„)
            if text.startswith('|') and '|' in text and len(text) > 10:
                # í…Œì´ë¸” ë¸”ë¡ì˜ ëª¨ë“  ë‹¨ë½ ìˆ˜ì§‘
                table_paragraphs = [paragraph]
                table_start_idx = i
                j = i + 1

                # ì—°ì†ëœ í…Œì´ë¸” í–‰ë“¤ ì°¾ê¸°
                while j < len(doc.paragraphs):
                    next_para = doc.paragraphs[j]
                    next_text = next_para.text.strip()

                    if next_text.startswith('|') and '|' in next_text:
                        table_paragraphs.append(next_para)
                        j += 1
                    else:
                        break

                # ì´ ìœ„ì¹˜ì— í…Œì´ë¸” ì‚½ì… ì˜ˆì •
                if table_paragraphs and markdown_tables:
                    table_insert_positions.append({
                        'start_para': table_paragraphs[0],
                        'paragraphs': table_paragraphs,
                        'markdown_table': markdown_tables.pop(0)
                    })

                i = j
            else:
                i += 1

        # ì—­ìˆœìœ¼ë¡œ ì²˜ë¦¬ (ì¸ë±ìŠ¤ ë³€ê²½ ë°©ì§€)
        for pos_info in reversed(table_insert_positions):
            start_para = pos_info['start_para']
            paragraphs = pos_info['paragraphs']
            md_table = pos_info['markdown_table']

            # í…Œì´ë¸” íŒŒì‹±
            rows_data = self._parse_markdown_table(md_table)
            if not rows_data:
                continue

            num_rows = len(rows_data)
            num_cols = max(len(row) for row in rows_data)

            # ê° í–‰ì˜ ì—´ ê°œìˆ˜ë¥¼ ë§ì¶¤
            for row in rows_data:
                while len(row) < num_cols:
                    row.append('')

            # ì²« ë²ˆì§¸ ë‹¨ë½ ìœ„ì¹˜ì— í…Œì´ë¸” ì‚½ì…
            p_element = start_para._element
            parent = p_element.getparent()

            # ìƒˆ í…Œì´ë¸” ìƒì„± (ë‹¨ë½ ì•ì— ì‚½ì…)
            from docx.table import Table
            tbl = doc.add_table(rows=num_rows, cols=num_cols)._element

            # ë‹¨ë½ ë°”ë¡œ ì•ì— í…Œì´ë¸” ì‚½ì…
            parent.insert(parent.index(p_element), tbl)

            # í…Œì´ë¸” ê°ì²´ ê°€ì ¸ì˜¤ê¸°
            new_table = Table(tbl, doc)

            try:
                new_table.style = 'Light Grid Accent 1'
            except:
                try:
                    new_table.style = 'Table Grid'
                except:
                    pass

            # ë°ì´í„° ì±„ìš°ê¸°
            for i, row_data in enumerate(rows_data):
                for j, cell_data in enumerate(row_data):
                    if j < num_cols and i < num_rows:
                        cell = new_table.rows[i].cells[j]
                        cell.text = str(cell_data) if cell_data else ''

                        # ì²« í–‰ì€ í—¤ë”ë¡œ ë³¼ë“œ ì²˜ë¦¬
                        if i == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.size = Pt(10)
                                    run.font.name = 'Malgun Gothic'
                        else:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                                    run.font.name = 'Malgun Gothic'

            # í…Œì´ë¸”ì— í•´ë‹¹í•˜ëŠ” í…ìŠ¤íŠ¸ ë‹¨ë½ë“¤ ì‚­ì œ
            for para in paragraphs:
                p = para._element
                p.getparent().remove(p)

    def _generate_word_with_pandoc(self, report_data: Dict[str, Any], output_path: str):
        """Pandocì„ ì‚¬ìš©í•œ Word ìƒì„± (ë§ˆí¬ë‹¤ìš´ ì™„ë²½ ì§€ì›)"""
        # ê²°ê³¼ ìˆ˜ì§‘
        results = report_data.get('results', [])
        markdown_content = []

        # ë³´ê³ ì„œ ì œëª© ì¶”ê°€ (ì§ˆë¬¸ ê¸°ë°˜)
        if results and results[0].get('question'):
            title = self._generate_report_title(results[0]['question'])
            markdown_content.append(f'# {title}\n\n')

            # ì‘ì„±ì ë° ì‘ì„±ì¼ì ì¶”ê°€
            author = report_data.get('author', 'Unknown')
            created_date = report_data.get('created_date', datetime.now().strftime("%Y-%m-%d"))
            markdown_content.append(f'**ì‘ì„±ì:** {author}  |  **ì‘ì„±ì¼:** {created_date}\n\n')
            markdown_content.append('---\n\n')

        for result in results:
            if result.get('success'):
                answer = result.get('answer', 'N/A')
                # í…Œì´ë¸” í˜•ì‹ ìˆ˜ì •
                answer = self._fix_table_format(answer)
                # "ì„ì›ë³´ê³ ì„œ" í—¤ë”© ì œê±° (ì²« ë²ˆì§¸ # í—¤ë”© ì œê±°)
                answer = self._remove_first_heading(answer)
                markdown_content.append(answer)
                markdown_content.append('\n\n')  # ì§ˆë¬¸ ì‚¬ì´ ê°„ê²©

        full_markdown = '\n'.join(markdown_content)

        # ì„ì‹œ ë§ˆí¬ë‹¤ìš´ íŒŒì¼ ìƒì„±
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp:
            tmp.write(full_markdown)
            tmp_path = tmp.name

        # ë””ë²„ê¹…: ë§ˆí¬ë‹¤ìš´ íŒŒì¼ ê²½ë¡œ ì¶œë ¥
        print(f"ğŸ“ ì„ì‹œ ë§ˆí¬ë‹¤ìš´ íŒŒì¼: {tmp_path}")

        try:
            # Pandocìœ¼ë¡œ ë³€í™˜
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)

            pypandoc.convert_file(
                tmp_path,
                'docx',
                outputfile=str(output_file),
                extra_args=['--reference-doc='] if False else []  # í•„ìš”ì‹œ í…œí”Œë¦¿ ì¶”ê°€
            )

            print(f"ğŸ“„ Word ë¬¸ì„œ ìƒì„± ì™„ë£Œ (Pandoc): {output_file}")

        finally:
            # ì„ì‹œ íŒŒì¼ ì‚­ì œ (ë””ë²„ê¹… ì‹œ ì£¼ì„ ì²˜ë¦¬)
            # Path(tmp_path).unlink(missing_ok=True)
            pass

    def _generate_word_basic(self, report_data: Dict[str, Any], output_path: str):
        """ê¸°ë³¸ ë°©ì‹ìœ¼ë¡œ Word ìƒì„± (python-docx)"""
        doc = Document()

        # ì²« ë²ˆì§¸ í—¤ë”© í”Œë˜ê·¸ ì´ˆê¸°í™”
        self.first_heading_added = False

        # ë³´ê³ ì„œ ì œëª© ì¶”ê°€ (ì§ˆë¬¸ ê¸°ë°˜)
        results = report_data.get('results', [])
        if results and results[0].get('question'):
            title = self._generate_report_title(results[0]['question'])
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.first_heading_added = True

            # ì‘ì„±ì ë° ì‘ì„±ì¼ì ì¶”ê°€
            author = report_data.get('author', 'Unknown')
            created_date = report_data.get('created_date', datetime.now().strftime("%Y-%m-%d"))

            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info_para.add_run(f"ì‘ì„±ì: {author}  |  ì‘ì„±ì¼: {created_date}")
            info_run.font.size = Pt(11)
            info_run.font.color.rgb = RGBColor(100, 100, 100)

            doc.add_paragraph()  # ì œëª© ë’¤ ê°„ê²©

        # ê²°ê³¼

        for i, result in enumerate(results, 1):
            if result.get('success'):
                # ë‹µë³€ë§Œ í‘œì‹œ (ì§ˆë¬¸ê³¼ ê²€ìƒ‰ëœ ë¬¸ì„œ ì •ë³´ëŠ” ì œì™¸)
                answer = result.get('answer', 'N/A')
                # í…Œì´ë¸” í˜•ì‹ ìˆ˜ì •
                answer = self._fix_table_format(answer)

                # ë§ˆí¬ë‹¤ìš´ í˜•ì‹ ì²˜ë¦¬
                self._add_formatted_text(doc, answer)

                # ì´ë¯¸ì§€ ì²¨ë¶€ (í•„í„°ë§ ì ìš©)
                images = result.get('images', [])
                if images:
                    # ê´€ë ¨ì„± ìˆëŠ” ì´ë¯¸ì§€ë§Œ í•„í„°ë§
                    relevant_images = [
                        img for img in images
                        if self._should_include_image(img, answer)
                    ]

                    if relevant_images:
                        doc.add_paragraph()  # ë‹µë³€ê³¼ ì´ë¯¸ì§€ ì‚¬ì´ ê°„ê²©

                        # ì´ë¯¸ì§€ ì„¹ì…˜ ì œëª©
                        para = doc.add_paragraph()
                        run = para.add_run("ğŸ“Š í•µì‹¬ ê·¸ë˜í”„ ë° ê²°ê³¼")
                        run.font.size = Pt(12)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(70, 70, 70)

                        # ê° ì´ë¯¸ì§€ ì‚½ì…
                        for img in relevant_images:
                            img_path = img.get('path')
                            img_desc = img.get('description')
                            img_source = img.get('source')

                            if img_path:
                                # ì´ë¯¸ì§€ ì„¤ëª… ì§§ê²Œ ìš”ì•½
                                caption = self._shorten_image_caption(img_desc) if img_desc else "ì´ë¯¸ì§€"
                                if img_source:
                                    caption += f" (ì¶œì²˜: {img_source})"

                                self._add_image(doc, img_path, caption, max_width=5.0)
                                doc.add_paragraph()  # ì´ë¯¸ì§€ ì‚¬ì´ ê°„ê²©

            else:
                # ì˜¤ë¥˜ ë°œìƒ
                self._add_paragraph(doc, f"âŒ ì˜¤ë¥˜ ë°œìƒ: {result.get('error', 'Unknown error')}", bold=True)

            # ì§ˆë¬¸ ì‚¬ì´ ê°„ê²©
            if i < len(results):
                doc.add_paragraph()
                doc.add_paragraph()

        # ì €ì¥
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_file)

        print(f"ğŸ“„ Word ë¬¸ì„œ ìƒì„± ì™„ë£Œ (ê¸°ë³¸): {output_file}")

    def generate_pdf_report(self, report_data: Dict[str, Any], output_path: str):
        """PDF ë³´ê³ ì„œ ìƒì„± (Wordë¥¼ ë¨¼ì € ë§Œë“¤ê³  PDFë¡œ ë³€í™˜)

        Args:
            report_data: ë³´ê³ ì„œ ë°ì´í„° (JSON)
            output_path: ì¶œë ¥ íŒŒì¼ ê²½ë¡œ
        """
        # ì„ì‹œ Word íŒŒì¼ ìƒì„±
        temp_docx = output_path.replace('.pdf', '_temp.docx')
        self.generate_word_report(report_data, temp_docx)

        try:
            # Wordë¥¼ PDFë¡œ ë³€í™˜ (LibreOffice ì‚¬ìš©)
            import subprocess

            output_file = Path(output_path)
            output_dir = output_file.parent

            # LibreOfficeë¡œ ë³€í™˜
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_dir),
                temp_docx
            ]

            result = subprocess.run(cmd, capture_output=True, text=True)

            if result.returncode == 0:
                # ìƒì„±ëœ PDF íŒŒì¼ì„ ì›í•˜ëŠ” ì´ë¦„ìœ¼ë¡œ ë³€ê²½
                generated_pdf = temp_docx.replace('.docx', '.pdf')
                if Path(generated_pdf).exists() and generated_pdf != output_path:
                    Path(generated_pdf).rename(output_path)

                print(f"ğŸ“„ PDF ë¬¸ì„œ ìƒì„± ì™„ë£Œ: {output_path}")
            else:
                print(f"âš ï¸ PDF ë³€í™˜ ì‹¤íŒ¨: {result.stderr}")
                print(f"ğŸ’¡ Word íŒŒì¼ì„ ì‚¬ìš©í•˜ì„¸ìš”: {temp_docx}")

        except Exception as e:
            print(f"âš ï¸ PDF ìƒì„± ì‹¤íŒ¨: {e}")
            print(f"ğŸ’¡ Word íŒŒì¼ì„ ì‚¬ìš©í•˜ì„¸ìš”: {temp_docx}")

        finally:
            # ì„ì‹œ Word íŒŒì¼ ì‚­ì œ (ì˜µì…˜)
            # Path(temp_docx).unlink(missing_ok=True)
            pass


def main():
    """í…ŒìŠ¤íŠ¸ìš© ë©”ì¸ í•¨ìˆ˜"""
    import json
    import argparse

    parser = argparse.ArgumentParser(description="ë³´ê³ ì„œ ë¬¸ì„œ ìƒì„±ê¸°")
    parser.add_argument("--json", type=str, required=True, help="ì…ë ¥ JSON íŒŒì¼")
    parser.add_argument("--output", type=str, required=True, help="ì¶œë ¥ íŒŒì¼ (.docx ë˜ëŠ” .pdf)")

    args = parser.parse_args()

    # JSON ë¡œë“œ
    with open(args.json, 'r', encoding='utf-8') as f:
        report_data = json.load(f)

    # ë¬¸ì„œ ìƒì„±
    generator = DocumentGenerator()

    if args.output.endswith('.pdf'):
        generator.generate_pdf_report(report_data, args.output)
    elif args.output.endswith('.docx'):
        generator.generate_word_report(report_data, args.output)
    else:
        print("âŒ ì§€ì›ë˜ì§€ ì•ŠëŠ” íŒŒì¼ í˜•ì‹ì…ë‹ˆë‹¤. .docx ë˜ëŠ” .pdfë¥¼ ì‚¬ìš©í•˜ì„¸ìš”.")


if __name__ == "__main__":
    main()
